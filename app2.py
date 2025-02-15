import time
import threading
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
import tkinter as tk
from tkinter import scrolledtext, messagebox

# Función de scraping para Hotfrog
def scrape_hotfrog(profesion):
    url = f"https://www.hotfrog.es/search/es/{profesion}"
    
    try:
        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
        driver.get(url)
        time.sleep(3)  # Esperar a que cargue la página
        
        resultados = []
        empresas = driver.find_elements(By.CLASS_NAME, "hf-box")

        for empresa in empresas:
            try:
                # Intentar extraer el nombre de la empresa
                nombre_elemento = empresa.find_element(By.XPATH, ".//h3/a/strong")
                nombre_empresa = nombre_elemento.text.strip()
            except:
                nombre_empresa = None  # En lugar de cadena vacía, ponemos None para filtrar después

            try:
                # Extraer el teléfono desde data-visible-number
                telefono_elemento = empresa.find_element(By.XPATH, ".//div[@class='w-100 small text-nowrap']/a")
                telefono = telefono_elemento.get_attribute("data-visible-number").strip()
            except:
                telefono = None  # Evita añadir números sin empresa

            # Solo agregar si tenemos nombre y teléfono
            if nombre_empresa and telefono:
                resultados.append(f"{nombre_empresa} - {telefono}")
        
        driver.quit()

        # Eliminar duplicados
        resultados = list(set(resultados))
        
        return resultados if resultados else ["No se encontraron resultados."]
    
    except Exception as e:
        return [f"Error al obtener datos: {e}"]

# Función para guardar en Word
def guardar_en_word(profesion, datos_hotfrog):
    doc = Document()
    doc.add_heading(f"Resultados de búsqueda para {profesion}", level=1)
    
    doc.add_heading("Hotfrog:", level=2)
    for anuncio in datos_hotfrog:
        doc.add_paragraph(f"- {anuncio}")
    
    archivo_nombre = f"resultados_{profesion}.docx"
    doc.save(archivo_nombre)
    messagebox.showinfo("Archivo guardado", f"Los resultados se han guardado en {archivo_nombre}")

# Interfaz gráfica con Tkinter
root = tk.Tk()
root.title("Scraper Hotfrog")
root.geometry("700x600")

tk.Label(root, text="Introduce una profesión:").pack(pady=5)
entrada = tk.Entry(root, width=60)
entrada.pack(pady=5)

def buscar():
    profesion = entrada.get()
    datos_hotfrog = scrape_hotfrog(profesion)
    threading.Thread(target=lambda: guardar_en_word(profesion, datos_hotfrog), daemon=True).start()

boton_buscar = tk.Button(root, text="Buscar", command=buscar)
boton_buscar.pack(pady=5)

resultado_text = scrolledtext.ScrolledText(root, width=80, height=15, state=tk.DISABLED)
resultado_text.pack(pady=5)

root.mainloop()
